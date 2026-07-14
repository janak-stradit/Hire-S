import hashlib
import os
import random
import re
import uuid
from collections import defaultdict
from copy import copy
from datetime import datetime, timedelta
from pathlib import Path

import pdfplumber
from docx import Document
from openpyxl import load_workbook
from PyPDF2 import PdfReader

from generate_hirex_1000_candidates import (
    BASE_DATE,
    CITIES_STATES,
    COMPANIES,
    DOMAIN_BLUEPRINTS,
    FINAL_COLUMNS,
    salary_by_experience,
)


ROOT = Path(__file__).resolve().parents[1]
RESUME_DIR = ROOT / "Resumes"
BASE_WORKBOOK_PATH = ROOT / "Reports" / "hirex_1000_candidate_dummy_data.xlsx"
WORKBOOK_PATH = ROOT / "Reports" / "hirex_candidate_master_with_extracted_resumes.xlsx"
TEMP_WORKBOOK_PATH = ROOT / "Reports" / ".hirex_candidate_import_tmp.xlsx"
IMPORT_TIME = datetime(2026, 6, 19, 12, 0, 0)

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

MANUAL_SCANNED_TEXT = {
    "Updated Resume Dhananjay.pdf": """Dhananjay Vilas Pandit
Phone: 919579446034
Email: dhananjaypandit9579@gmail.com
Location: Yawal, Maharashtra, 425301

ABOUT ME
Motivated BCA graduate from Yawal, Jalgaon with strong teamwork, time management, and a positive attitude. Eager to apply my academic knowledge in a professional environment and grow within the IT industry.

TECHNICAL SKILLS
Programming Languages: C, C++, Core Java
Basic HTML/CSS understanding
MS Office: Word, Excel, PowerPoint
Internet Research and Data Handling

SOFT SKILLS
Teamwork
Time Management
Positive Attitude
Communication Skills
Adaptability
Problem-Solving

EDUCATION
Bachelor of Computer Applications (BCA)
Kavayitri Bahinabai Chaudhari North Maharashtra University, Jalgaon
2023-2025
TYBCA: 68.19% Year 2025
SYBCA: 65% Year 2024
FYBCA: 65.86% Year 2023

Higher Secondary Certificate (HSC)
Sane Guruji Madhyamik Vidyalaya, Yawal
2019-2022
Percentage: 73%

Secondary School Certificate (SSC)
Bal Sanskar Madhyamik Vidya, Yawal
2019-2020
Percentage: 76%

EXPERIENCE
Internship - CodingSamurai Online Platform
December 2024-January 2025
- Developed and optimized C++ programs, improving efficiency and performance.
- Implemented data structures and algorithms to solve real-world problems.
- Gained hands-on experience in object-oriented programming and memory management.
- Strengthened understanding of file handling and multithreading concepts.

PROJECT
Online Gas Booking System - April 2025
Developed a user-friendly web application for booking LPG gas cylinders online. The system supports new connection requests, booking management, transaction history, and admin control for monitoring and updating user records. It improves service efficiency by digitizing the entire gas booking workflow.

TECHNOLOGIES USED
Frontend: HTML, CSS, JavaScript
Backend: PHP
Database: MySQL
Tools and Environment: XAMPP, phpMyAdmin, VS Code, Apache Server

STRENGTHS
Teamwork, Time Management, Positive Attitude, Quick Learner

LANGUAGES
English, Hindi, Marathi""",
}

SCANNED_DUPLICATE_ALIASES = {
    "Copy of Beige Dark Gray Minimalist Web Developer Resume.pdf": "ravindranathvajire@gmail.com",
    "Rohit_ingle_documents.pdf": "rohitingle96424@gmail.com",
    "Rohit_Ingle_Resume_45.pdf": "rohitingle96424@gmail.com",
}

EMAIL_EXCLUSIONS = ("helpdesk", "support@", "payments.", "noreply", "example.com")

DEGREE_PATTERNS = [
    (r"\bMaster of Computer Applications\b|\bM\.?C\.?A\.?\b", "Master of Computer Applications (MCA)"),
    (r"\bMaster of Science\b|\bM\.?S\.?c\.?\b", "Master of Science (M.Sc.)"),
    (r"\bMaster of Technology\b|\bM\.?Tech\.?\b", "Master of Technology (M.Tech.)"),
    (r"\bMaster of Business Administration\b|\bMBA\b", "Master of Business Administration (MBA)"),
    (r"\bMaster of Commerce\b|\bM\.?Com\.?\b", "Master of Commerce (M.Com.)"),
    (r"\bBachelor of Technology\b|\bB\.?Tech\.?\b", "Bachelor of Technology (B.Tech.)"),
    (r"\bBachelor of Engineering\b|\bB\.?E\.?\b", "Bachelor of Engineering (B.E.)"),
    (r"\bBachelor of Computer Applications\b|\bB\.?C\.?A\.?\b", "Bachelor of Computer Applications (BCA)"),
    (r"\bBachelor of Science\b|\bB\.?S\.?c\.?\b", "Bachelor of Science (B.Sc.)"),
    (r"\bBachelor of Commerce\b|\bB\.?Com\.?\b", "Bachelor of Commerce (B.Com.)"),
    (r"\bDiploma in Mechanical Engineering\b", "Diploma in Mechanical Engineering"),
    (r"\bDiploma in Information Technology\b", "Diploma in Information Technology"),
    (r"\bCA Intermediate\b|\bCA Inter\b", "CA Intermediate"),
]

ROLE_PATTERNS = [
    "SAP CO Consultant", "SAP FICO Consultant", "Identity Specialist", "Identity Services Specialist",
    "Network Documentation Specialist", "Network Architect", "Shift Engineer Mechanical",
    "Mechanical Engineer", "Java Backend Developer", ".NET Developer", "Full Stack Developer",
    "React Developer", "Python Developer", "Software Engineer", "Data Analyst", "Data Engineer",
    "DevOps Engineer", "System Administrator", "Senior Quality Analyst", "QA Engineer",
    "Workforce Management Specialist", "Technical Support Engineer", "Security Engineer",
    "Business Analyst", "Trainee Engineer",
]

SKILL_TERMS = sorted(
    {
        skill
        for blueprint in DOMAIN_BLUEPRINTS.values()
        for skills in blueprint["skills"].values()
        for skill in skills
    }
    | {
        "ABAP", "SAP CO", "SAP FICO", "SAP S/4HANA", "SAP ECC", "Cost Center Accounting",
        "Profit Center Accounting", "Product Costing", "Active Directory", "Azure AD", "Entra ID",
        "Identity and Access Management", "IAM", "SailPoint", "CyberArk", "Okta", "PowerShell",
        "Cisco", "BGP", "OSPF", "MPLS", "VLAN", "Routing", "Switching", "Network Documentation",
        "AutoCAD", "SolidWorks", "Mechanical Maintenance", "Preventive Maintenance", "PLC", "CNC",
        "C", "C++", "Core Java", "JavaScript", "HTML", "CSS", "PHP", "MySQL", ".NET", "C#",
        "Node.js", "Angular", "MongoDB", "Spring", "Hibernate", "Microservices", "Bootstrap",
    },
    key=len,
    reverse=True,
)

SECTION_HEADINGS = {
    "summary": ["professional summary", "profile summary", "career summary", "summary", "about me", "objective", "career objective"],
    "skills": ["technical skills", "core competencies", "key skills", "skills", "technical expertise", "competencies"],
    "education": ["education", "academic qualification", "academic qualifications", "academics", "educational qualification"],
    "certifications": ["certifications", "certificates", "professional certifications", "certification"],
    "projects": ["projects", "academic projects", "project details", "projects & dashboards", "project"],
    "experience": ["professional experience", "work experience", "experience", "employment history", "career history"],
}

ALL_HEADINGS = sorted({heading for headings in SECTION_HEADINGS.values() for heading in headings}, key=len, reverse=True)

NAME_OVERRIDES = {
    "ATS RESUME1.pdf": "Thota Gopi Raju",
    "Java_Developer-1.pdf": "Rohit Ingle",
    "KiranKumar-Resume.pdf": "Kiran Kumar Mankali",
    "Mahantesh_Ganjagol_CV.pdf": "Mahantesh Ganjagol",
    "Monica Resume .pdf": "Monica Vijaykumar",
    "Purple and White Clean and Professional Resume.pdf": "Manish Shirsat",
    "Ravi_Kumar_Resume.pdf": "Dasa Ravi Kumar",
    "Ravi_Resume_2   PDF.pdf": "Ravindranath Vajire",
    "Ravi_Resume_2 (1)-1.pdf": "Ravindranath Vajire",
    "Pramodyadav Shift Engineer-GHRS.docx": "Pramod Yadav",
    "RashiJain_SAP_CO_Consultant _GHRS.pdf": "Rashi M Jain",
    "resume.pdf": "Uday Rajendra Sapkal",
    "resume-2.pdf": "Yash Chatse",
    "resume6.pdf": "Manish Shirsat",
    "ShaikhSalimHanif_Net Developer_GHRS (1).pdf": "Salim Shaikh",
    "Updated Resume Dhananjay.pdf": "Dhananjay Vilas Pandit",
}

DOMAIN_OVERRIDES = {
    "Parikshit sharma.pdf": "HR",
    "Pooja Bhandare.pdf": "Testing",
    "Rayavaram gangadhara_Trainee Streling_GHRS.pdf": "IT Support",
    "Shriprasad Wable_Network Documentation Specialist _GHRS.pdf": "Cybersecurity",
}

ROLE_OVERRIDES = {
    "Parikshit sharma.pdf": "Workforce Management Specialist",
    "Pooja Bhandare.pdf": "Senior Quality Analyst",
    "Rayavaram gangadhara_Trainee Streling_GHRS.pdf": "Technical Support Engineer",
    "Shriprasad Wable_Network Documentation Specialist _GHRS.pdf": "Security Engineer",
}

GENERIC_FILENAME_PREFIXES = {
    "ats", "copy", "java", "purple", "resume", "updated",
}


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]+", "\n", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def extract_document_text(path: Path) -> str:
    if path.name in MANUAL_SCANNED_TEXT:
        return clean_text(MANUAL_SCANNED_TEXT[path.name])
    if path.suffix.lower() == ".docx":
        document = Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return clean_text("\n".join(parts))

    text = ""
    try:
        with pdfplumber.open(path) as pdf:
            text = "\n".join(
                page.extract_text(x_tolerance=2, y_tolerance=3, layout=False) or "" for page in pdf.pages
            )
    except Exception:
        text = ""
    if len(text.strip()) < 100:
        try:
            text = "\n".join(page.extract_text() or "" for page in PdfReader(path).pages)
        except Exception:
            pass
    return clean_text(text)


def extract_emails(text: str) -> list[str]:
    values = re.findall(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    unique = []
    for value in values:
        email = value.strip(".,;:").lower()
        if email not in unique:
            unique.append(email)
    return unique


def select_personal_email(emails: list[str]) -> str | None:
    for email in emails:
        if not any(excluded in email for excluded in EMAIL_EXCLUSIONS):
            return email
    return emails[0] if emails else None


def extract_phones(text: str) -> list[str]:
    candidates = re.findall(r"(?:\+?91[\s-]?)?[6-9](?:[\s()-]?\d){9}", text)
    unique = []
    for candidate in candidates:
        digits = re.sub(r"\D", "", candidate)
        if len(digits) == 12 and digits.startswith("91"):
            digits = digits[2:]
        if len(digits) == 10 and digits not in unique:
            unique.append(digits)
    return unique


def filename_name(path: Path) -> str:
    stem = path.stem
    stem = re.split(r"[_-](?:Resume|Updated|SAP|Identity|Network|Shift|Java|Dotnet|Net Developer|FullStack|React|Trainee)", stem, flags=re.I)[0]
    stem = re.sub(r"\b(?:resume|updated|cv|ghrs|ats|pdf|copy of|documents?)\b", " ", stem, flags=re.I)
    stem = re.sub(r"\(\d+\)|\d{8,}", " ", stem)
    stem = re.sub(r"(?<=[a-z])(?=[A-Z])", " ", stem)
    return re.sub(r"[_-]+|\s{2,}", " ", stem).strip(" ._-")


def extract_name(text: str, path: Path) -> str:
    if path.name in NAME_OVERRIDES:
        return NAME_OVERRIDES[path.name]
    file_candidate = filename_name(path)
    file_words = file_candidate.split()
    if (
        2 <= len(file_words) <= 5
        and file_words[0].lower() not in GENERIC_FILENAME_PREFIXES
        and not any(char.isdigit() for char in file_candidate)
    ):
        return file_candidate.title()

    ignored = {
        "resume", "curriculum vitae", "professional summary", "profile summary", "summary", "objective",
        "contact", "education", "skills", "technical skills", "experience", "work experience",
    }
    lines = [re.sub(r"\s+", " ", line).strip(" |:-") for line in text.splitlines() if line.strip()]
    for line in lines[:20]:
        lowered = line.lower()
        words = re.findall(r"[A-Za-z]+(?:\.[A-Za-z]+)?", line)
        if (
            lowered not in ignored
            and 2 <= len(words) <= 5
            and "@" not in line
            and not re.search(r"\d", line)
            and not any(term in lowered for term in [
                "developer", "engineer", "consultant", "specialist", "analyst", "manager",
                "address", "reservation", "ticket", "efficiency", "management", "access",
                "skills", "github", "linkedin", "computer science student",
            ])
        ):
            return " ".join(word.capitalize() if not word.isupper() else word.title() for word in words)
    return file_candidate.title() if file_candidate else f"Candidate {path.stem[:20]}"


def split_name(full_name: str) -> tuple[str, str]:
    parts = [part for part in full_name.split() if part]
    if not parts:
        return "Candidate", "Unknown"
    if len(parts) == 1:
        return parts[0], "Candidate"
    return " ".join(parts[:-1]), parts[-1]


def extract_section(text: str, section: str) -> str:
    heading_pattern = "|".join(re.escape(item) for item in SECTION_HEADINGS[section])
    next_pattern = "|".join(re.escape(item) for item in ALL_HEADINGS)
    pattern = re.compile(
        rf"(?ims)^\s*(?:{heading_pattern})\s*[:\-]?\s*$\s*(.*?)(?=^\s*(?:{next_pattern})\s*[:\-]?\s*$|\Z)"
    )
    match = pattern.search(text)
    if not match:
        return ""
    return clean_text(match.group(1))[:12000]


def extract_links(text: str) -> tuple[str, str, str]:
    urls = re.findall(r"(?:https?://|www\.)[^\s<>]+", text, flags=re.I)
    linkedin = next((url.strip(".,);]") for url in urls if "linkedin" in url.lower()), "")
    github = next((url.strip(".,);]") for url in urls if "github" in url.lower()), "")
    portfolio = next(
        (url.strip(".,);]") for url in urls if "linkedin" not in url.lower() and "github" not in url.lower()),
        "",
    )
    return linkedin, github, portfolio


def classify_domain(text: str, path: Path) -> str:
    if path.name in DOMAIN_OVERRIDES:
        return DOMAIN_OVERRIDES[path.name]
    source = f"{path.stem} {text[:10000]}".lower()
    rules = [
        ("SAP Finance and Controlling", ["sap co", "sap fico", "product costing", "cost center accounting"]),
        ("Identity and Access Management", ["identity specialist", "identity services", "iam", "sailpoint", "cyberark", "azure ad"]),
        ("Network Engineering", ["network architect", "network documentation", "bgp", "ospf", "mpls"]),
        ("Mechanical Engineering", ["shift engineer", "mechanical engineer", "preventive maintenance", "solidworks", "autocad"]),
        ("Data Analytics", ["data analyst", "power bi", "tableau", "data visualization"]),
        ("Data Engineering", ["data engineer", "pyspark", "airflow", "databricks"]),
        ("DevOps", ["devops", "kubernetes", "terraform", "jenkins"]),
        ("Testing", ["qa engineer", "software testing", "selenium", "manual testing"]),
        ("Software Development", ["developer", "java", "react", "full stack", ".net", "python", "spring boot"]),
        ("Finance", ["financial analyst", "investment", "budgeting", "forecasting"]),
        ("HR", ["human resources", "talent acquisition", "recruiter", "payroll"]),
        ("Accounting", ["accountant", "gst", "taxation", "ledger"]),
        ("Sales and Marketing", ["sales", "marketing", "seo", "lead generation"]),
        ("IT Support", ["it support", "desktop support", "service desk", "system administrator"]),
    ]
    scores = [(domain, sum(source.count(keyword) for keyword in keywords)) for domain, keywords in rules]
    best_domain, score = max(scores, key=lambda item: item[1])
    return best_domain if score else "Software Development"


def extract_role(text: str, path: Path, domain: str) -> str:
    if path.name in ROLE_OVERRIDES:
        return ROLE_OVERRIDES[path.name]
    source = f"{path.stem}\n{text[:8000]}"
    role_domains = {
        "SAP CO Consultant": "SAP Finance and Controlling",
        "SAP FICO Consultant": "SAP Finance and Controlling",
        "Identity Specialist": "Identity and Access Management",
        "Identity Services Specialist": "Identity and Access Management",
        "Network Documentation Specialist": "Network Engineering",
        "Network Architect": "Network Engineering",
        "Shift Engineer Mechanical": "Mechanical Engineering",
        "Mechanical Engineer": "Mechanical Engineering",
        "Java Backend Developer": "Software Development",
        ".NET Developer": "Software Development",
        "Full Stack Developer": "Software Development",
        "React Developer": "Software Development",
        "Python Developer": "Software Development",
        "Software Engineer": "Software Development",
        "Data Analyst": "Data Analytics",
        "Data Engineer": "Data Engineering",
        "DevOps Engineer": "DevOps",
        "Senior Quality Analyst": "Testing",
        "QA Engineer": "Testing",
        "Workforce Management Specialist": "HR",
        "Technical Support Engineer": "IT Support",
        "Security Engineer": "Cybersecurity",
        "Business Analyst": "Data Analytics",
        "Trainee Engineer": "Mechanical Engineering",
    }
    for role in ROLE_PATTERNS:
        if re.search(re.escape(role), source, flags=re.I) and role_domains.get(role, domain) == domain:
            return role
    domain_defaults = {
        "SAP Finance and Controlling": "SAP CO Consultant",
        "Identity and Access Management": "Identity Specialist",
        "Network Engineering": "Network Engineer",
        "Mechanical Engineering": "Mechanical Engineer",
        "Software Development": "Software Developer",
        "Data Analytics": "Data Analyst",
        "Data Engineering": "Data Engineer",
        "DevOps": "DevOps Engineer",
        "Testing": "QA Engineer",
        "Finance": "Financial Analyst",
        "HR": "HR Executive",
        "Accounting": "Accountant",
        "Sales and Marketing": "Sales and Marketing Executive",
        "IT Support": "IT Support Engineer",
        "Cybersecurity": "Security Engineer",
    }
    return domain_defaults[domain]


def extract_experience(text: str) -> float:
    patterns = [
        r"(?:total\s+)?experience\s*(?:of|:|-)?\s*(\d+(?:\.\d+)?)\+?\s*years?",
        r"(\d+(?:\.\d+)?)\+?\s*years?\s+(?:of\s+)?(?:professional\s+)?experience",
        r"over\s+(\d+(?:\.\d+)?)\s*years?",
    ]
    values = []
    for pattern in patterns:
        values.extend(float(item) for item in re.findall(pattern, text, flags=re.I))
    if values:
        return round(min(max(values), 30.0), 1)

    date_ranges = re.findall(
        r"(?i)(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s,.'-]*(20\d{2})\s*(?:-|–|to)\s*(?:(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s,.'-]*(20\d{2})|present|current)",
        text,
    )
    months = 0
    month_index = {name.lower(): number for number, name in enumerate(["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"], 1)}
    for start_month, start_year, end_month, end_year in date_ranges:
        start = int(start_year) * 12 + month_index[start_month[:3].lower()]
        end = IMPORT_TIME.year * 12 + IMPORT_TIME.month if not end_year else int(end_year) * 12 + month_index[end_month[:3].lower()]
        months += max(0, end - start)
    return round(min(months / 12, 30.0), 1) if months else 0.5


def level_from_experience(experience: float) -> str:
    if experience < 1:
        return "Fresher"
    if experience < 3:
        return "Associate"
    if experience < 5:
        return "Executive"
    if experience < 8:
        return "Senior"
    if experience < 11:
        return "Lead"
    if experience < 15:
        return "Manager"
    return "Senior Manager"


def extract_highest_education(text: str, domain: str) -> tuple[str, bool]:
    for pattern, label in DEGREE_PATTERNS:
        match = re.search(pattern, text, flags=re.I)
        if match:
            return label, False
    defaults = {
        "Mechanical Engineering": "Diploma in Mechanical Engineering",
        "SAP Finance and Controlling": "Bachelor of Commerce",
        "Finance": "Bachelor of Commerce",
        "Accounting": "Bachelor of Commerce",
        "HR": "Bachelor's Degree",
    }
    return defaults.get(domain, "Bachelor's Degree"), True


def extract_skills(text: str) -> tuple[str, bool]:
    section = extract_section(text, "skills")
    if section:
        return section[:6000], False
    found = [skill for skill in SKILL_TERMS if re.search(rf"(?i)(?<!\w){re.escape(skill)}(?!\w)", text)]
    if found:
        return ", ".join(found[:40]), False
    return "[Synthetic completion] Communication, Documentation, MS Office, Problem Solving", True


def extract_company(text: str) -> tuple[str, bool]:
    experience = extract_section(text, "experience")
    source = experience or text
    for company in sorted(COMPANIES, key=len, reverse=True):
        if re.search(rf"(?i)\b{re.escape(str(company))}\b", source):
            return str(company), False
    patterns = [
        r"(?im)^\s*(?:company|organization|employer)\s*[:\-]\s*([^\n|]{2,100})",
        r"(?im)^\s*([A-Z][A-Za-z0-9&.,'() -]{2,80}(?:Ltd\.?|Limited|LLC|Pvt\.? Ltd\.?|Technologies|Solutions|Systems|Services))\s*$",
    ]
    for pattern in patterns:
        match = re.search(pattern, source)
        if match:
            value = re.sub(r"\s+", " ", match.group(1)).strip(" -|:.")
            lowered = value.lower()
            if any(lowered.startswith(term) for term in [
                "powered by", "built ", "developed ", "designed ", "supported ",
                "implemented ", "network security implementation",
            ]):
                if " at " in value:
                    value = value.rsplit(" at ", 1)[-1].strip()
                else:
                    continue
            value = re.sub(r"(?i)^organization name\s*", "", value).strip(" -|:.")
            if 2 <= len(value) <= 120:
                return value[:180], False
    return "Confidential Employer [Synthetic completion]", True


def is_resume_document(path: Path, text: str) -> bool:
    if path.name in MANUAL_SCANNED_TEXT or path.name in SCANNED_DUPLICATE_ALIASES:
        return True
    lowered = text.lower()
    filename_signal = bool(re.search(r"resume|cv|ghrs|developer|engineer|consultant|specialist|architect", path.stem, flags=re.I))
    content_signals = sum(
        bool(re.search(pattern, lowered))
        for pattern in [
            r"\bskills?\b", r"\beducation\b|\bacademic", r"\bexperience\b|\bemployment",
            r"\bprojects?\b", r"\bsummary\b|\bobjective\b|\bprofile\b",
        ]
    )
    return filename_signal or content_signals >= 2


def find_location(text: str, seed: str) -> tuple[str, str, str, bool]:
    states = {city.lower(): (city, state) for city, state in CITIES_STATES}
    states.update({"pune": ("Pune", "Maharashtra"), "yawal": ("Yawal", "Maharashtra"), "jalgaon": ("Jalgaon", "Maharashtra"), "chikodi": ("Chikodi", "Karnataka")})
    lowered = text.lower()
    for key, (city, state) in states.items():
        if re.search(rf"\b{re.escape(key)}\b", lowered):
            return city, state, "India", False
    city, state = CITIES_STATES[int(hashlib.sha256(seed.encode()).hexdigest()[:8], 16) % len(CITIES_STATES)]
    return city, state, "India", True


def synthetic_phone(seed: str, used_phones: set[str]) -> str:
    number = 7000000000 + int(hashlib.sha256(seed.encode()).hexdigest()[:12], 16) % 2999999999
    phone = str(number)
    while phone in used_phones:
        phone = str(int(phone) + 1)
    return phone


def synthetic_email(first_name: str, last_name: str, seed: str, used_emails: set[str]) -> str:
    slug = re.sub(r"[^a-z0-9]", "", f"{first_name}.{last_name}".lower()) or "candidate"
    suffix = hashlib.sha256(seed.encode()).hexdigest()[:8]
    email = f"{slug}.{suffix}@hirex-synthetic.example"
    counter = 1
    while email in used_emails:
        email = f"{slug}.{suffix}{counter}@hirex-synthetic.example"
        counter += 1
    return email


def deterministic_id(prefix: str, identity: str) -> str:
    return f"{prefix}_{uuid.uuid5(uuid.NAMESPACE_URL, f'hirex-real-{prefix}-{identity}').hex[:12]}"


def candidate_identity(path: Path, text: str) -> str:
    if path.name in SCANNED_DUPLICATE_ALIASES:
        return f"email:{SCANNED_DUPLICATE_ALIASES[path.name]}"
    email = select_personal_email(extract_emails(text))
    if email:
        return f"email:{email}"
    phones = extract_phones(text)
    if phones:
        return f"phone:{phones[0]}"
    return f"name:{re.sub(r'[^a-z0-9]', '', extract_name(text, path).lower())}"


def inventory_unique_resumes() -> tuple[list[dict], dict]:
    records = []
    file_hash_groups: dict[str, list[str]] = defaultdict(list)
    identity_groups: dict[str, list[dict]] = defaultdict(list)
    files = sorted(path for path in RESUME_DIR.iterdir() if path.suffix.lower() in SUPPORTED_EXTENSIONS)

    rejected_non_resumes = []
    for path in files:
        file_hash = hashlib.sha256(path.read_bytes()).hexdigest()
        file_hash_groups[file_hash].append(path.name)
        text = extract_document_text(path)
        if not is_resume_document(path, text):
            rejected_non_resumes.append(path.name)
            continue
        identity = candidate_identity(path, text)
        record = {"path": path, "text": text, "file_hash": file_hash, "identity": identity}
        identity_groups[identity].append(record)
        records.append(record)

    selected = []
    duplicate_versions = 0
    for group in identity_groups.values():
        usable = [record for record in group if len(record["text"]) >= 100]
        if not usable:
            continue
        richest = max(usable, key=lambda record: len(record["text"]))
        selected.append(richest)
        duplicate_versions += len(group) - 1

    stats = {
        "files_found": len(files),
        "pdf_files": sum(path.suffix.lower() == ".pdf" for path in files),
        "docx_files": sum(path.suffix.lower() == ".docx" for path in files),
        "unique_candidates_selected": len(selected),
        "duplicate_versions_removed": duplicate_versions,
        "manual_scanned_resumes_transcribed": sum(record["path"].name in MANUAL_SCANNED_TEXT for record in selected),
        "unusable_unique_groups": sum(not any(len(record["text"]) >= 100 for record in group) for group in identity_groups.values()),
        "rejected_non_resume_files": rejected_non_resumes,
    }
    return sorted(selected, key=lambda record: record["path"].name.lower()), stats


def map_resume_to_row(
    record: dict,
    used_emails: set[str],
    used_phones: set[str],
    used_ids: set[str],
) -> tuple[dict, dict]:
    path: Path = record["path"]
    text: str = record["text"][:32767]
    identity: str = record["identity"]
    full_name = extract_name(text, path)
    first_name, last_name = split_name(full_name)

    emails = extract_emails(text)
    email = select_personal_email(emails)
    email_synthetic = not email or email in used_emails
    if email_synthetic:
        email = synthetic_email(first_name, last_name, identity, used_emails)
    assert email is not None
    used_emails.add(email)

    phones = extract_phones(text)
    phone = next((item for item in phones if item not in used_phones), None)
    phone_synthetic = phone is None
    if phone_synthetic:
        phone = synthetic_phone(identity, used_phones)
    used_phones.add(phone)

    domain = classify_domain(text, path)
    role = extract_role(text, path, domain)
    experience = extract_experience(text)
    level = level_from_experience(experience)
    company, company_synthetic = extract_company(text)
    highest_education, education_level_synthetic = extract_highest_education(text, domain)
    skills, skills_synthetic = extract_skills(text)
    city, state, country, location_synthetic = find_location(text, identity)
    linkedin, github, portfolio = extract_links(text)

    summary = extract_section(text, "summary")
    summary_synthetic = not bool(summary)
    if summary_synthetic:
        summary = (
            f"[Synthetic completion] {role} with {experience:.1f} years of aligned experience in {domain}. "
            f"Profile synthesized from the skills, education, projects, and employment evidence available in the source resume."
        )

    education = extract_section(text, "education")
    education_synthetic = not bool(education)
    if education_synthetic:
        education = f"[Synthetic completion] {highest_education}; institution and score not listed in the source resume."

    certifications = extract_section(text, "certifications")
    certifications_synthetic = not bool(certifications)
    if certifications_synthetic:
        certifications = "[Synthetic completion] No certification details listed; profile retained for agent verification."

    projects = extract_section(text, "projects")
    projects_synthetic = not bool(projects)
    if projects_synthetic:
        projects = f"[Synthetic completion] Role-aligned {domain} project evidence not explicitly listed; verify during screening."

    uploaded_at = datetime.fromtimestamp(path.stat().st_mtime).replace(microsecond=0)
    created_at = uploaded_at - timedelta(days=7)
    applied_at = uploaded_at + timedelta(days=1)
    if applied_at > IMPORT_TIME:
        applied_at = IMPORT_TIME

    candidate_id = deterministic_id("cand", identity)
    while candidate_id in used_ids:
        candidate_id = deterministic_id("cand", f"{identity}-{path.name}")
    used_ids.add(candidate_id)

    row = {
        "candidate_id": candidate_id,
        "resume_id": deterministic_id("res", identity),
        "parsed_resume_id": deterministic_id("parsed_resume", identity),
        "email": email,
        "password_hash": "$pbkdf2-sha256$synthetic$hirex_resume_import",
        "role": "candidate",
        "is_active": True,
        "created_at": created_at,
        "first_name": first_name,
        "last_name": last_name,
        "phone": phone,
        "city": city,
        "state": state,
        "country": country,
        "domain": domain,
        "level": level,
        "current_company": company,
        "current_role": role,
        "total_experience": experience,
        "expected_salary": salary_by_experience(experience),
        "notice_period": random.Random(identity).choice(["Immediate", "15 Days", "30 Days", "45 Days", "60 Days", "90 Days"]),
        "highest_education": highest_education,
        "linkedin_url": linkedin or f"https://linkedin.com/in/hirex-synthetic-{candidate_id[-8:]}",
        "github_url": github or f"https://github.com/hirex-synthetic-{candidate_id[-8:]}",
        "portfolio_url": portfolio or f"https://portfolio-{candidate_id[-8:]}.hirex-synthetic.example",
        "profile_completion_percentage": 100,
        "filename": path.name,
        "original_filename": path.name,
        "file_size": path.stat().st_size,
        "file_type": "application/pdf" if path.suffix.lower() == ".pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "storage_path": str(path.resolve()),
        "uploaded_at": uploaded_at,
        "skills": skills,
        "total_experience_years": experience,
        "education": education,
        "certifications": certifications,
        "projects": projects,
        "summary": summary,
        "raw_text": text,
        "parsed_at": IMPORT_TIME,
        "application_status": "Applied",
        "applied_at": applied_at,
        "resume": text,
    }
    provenance = {
        "email_synthetic": email_synthetic,
        "phone_synthetic": phone_synthetic,
        "location_synthetic": location_synthetic,
        "company_synthetic": company_synthetic,
        "highest_education_synthetic": education_level_synthetic,
        "skills_synthetic": skills_synthetic,
        "summary_synthetic": summary_synthetic,
        "education_synthetic": education_synthetic,
        "certifications_synthetic": certifications_synthetic,
        "projects_synthetic": projects_synthetic,
    }
    return row, provenance


def append_rows_to_workbook(rows: list[dict]) -> dict:
    workbook = load_workbook(BASE_WORKBOOK_PATH)
    worksheet = workbook["candidates"]
    headers = [cell.value for cell in worksheet[1]]
    if headers != FINAL_COLUMNS:
        raise ValueError("Workbook columns differ from the required 43-column contract")

    existing_rows = list(worksheet.iter_rows(min_row=2, values_only=True))
    email_index = headers.index("email")
    phone_index = headers.index("phone")
    id_index = headers.index("candidate_id")
    filename_index = headers.index("original_filename")
    used_emails = {str(row[email_index]).lower() for row in existing_rows if row[email_index]}
    used_phones = {str(row[phone_index]) for row in existing_rows if row[phone_index]}
    used_ids = {str(row[id_index]) for row in existing_rows if row[id_index]}
    imported_filenames = {str(row[filename_index]) for row in existing_rows if row[filename_index]}

    mapped_rows = []
    provenance_counts = defaultdict(int)
    skipped_already_imported = 0
    for record in rows:
        if record["path"].name in imported_filenames:
            skipped_already_imported += 1
            continue
        mapped, provenance = map_resume_to_row(record, used_emails, used_phones, used_ids)
        mapped_rows.append(mapped)
        for key, value in provenance.items():
            provenance_counts[key] += int(value)

    if mapped_rows:
        template_row = 2
        start_row = worksheet.max_row + 1
        for offset, mapped in enumerate(mapped_rows):
            target_row = start_row + offset
            for column, header in enumerate(headers, start=1):
                source_cell = worksheet.cell(row=template_row, column=column)
                target_cell = worksheet.cell(row=target_row, column=column, value=mapped[header])
                if source_cell.has_style:
                    target_cell.font = copy(source_cell.font)  # type: ignore
                    target_cell.border = copy(source_cell.border)  # type: ignore
                    target_cell.fill = copy(source_cell.fill)  # type: ignore
                    target_cell.protection = copy(source_cell.protection)  # type: ignore
                target_cell.number_format = source_cell.number_format
                target_cell.alignment = copy(source_cell.alignment)  # type: ignore
            worksheet.row_dimensions[target_row].height = worksheet.row_dimensions[template_row].height

        table = worksheet.tables.get("HireXCandidates")
        if table:
            table.ref = f"A1:AQ{worksheet.max_row}"
        worksheet.auto_filter.ref = f"A1:AQ{worksheet.max_row}"

    workbook.save(TEMP_WORKBOOK_PATH)
    workbook.close()

    validation = validate_workbook(TEMP_WORKBOOK_PATH, len(existing_rows), len(mapped_rows))
    os.replace(TEMP_WORKBOOK_PATH, WORKBOOK_PATH)
    return {
        "existing_rows": len(existing_rows),
        "rows_appended": len(mapped_rows),
        "skipped_already_imported": skipped_already_imported,
        "synthetic_field_counts": dict(provenance_counts),
        **validation,
    }


def validate_workbook(path: Path, existing_count: int, appended_count: int) -> dict:
    workbook = load_workbook(path, read_only=False, data_only=True)
    worksheet = workbook["candidates"]
    headers = [cell.value for cell in worksheet[1]]
    rows = list(worksheet.iter_rows(min_row=2, values_only=True))
    index = {name: headers.index(name) for name in headers}

    required_columns = [
        "candidate_id", "email", "phone", "first_name", "last_name", "domain", "current_role",
        "total_experience", "highest_education", "skills", "education", "certifications", "projects",
        "summary", "raw_text", "application_status", "applied_at", "resume",
    ]
    blank_required_cells = sum(
        value is None or (isinstance(value, str) and not value.strip())
        for row in rows
        for value in (row[index[column]] for column in required_columns)
    )
    duplicate_emails = len(rows) - len({str(row[index["email"]]).lower() for row in rows})
    duplicate_phones = len(rows) - len({str(row[index["phone"]]) for row in rows})
    duplicate_ids = len(rows) - len({str(row[index["candidate_id"]]) for row in rows})
    imported_rows = rows[-appended_count:] if appended_count else []
    raw_mismatches = sum(row[index["raw_text"]] != row[index["resume"]] for row in imported_rows)
    over_limit = sum(len(str(row[index["resume"]] or "")) > 32767 for row in imported_rows)

    checks = {
        "final_rows": len(rows),
        "final_columns": len(headers),
        "blank_required_cells": blank_required_cells,
        "duplicate_emails": duplicate_emails,
        "duplicate_phones": duplicate_phones,
        "duplicate_candidate_ids": duplicate_ids,
        "imported_resume_raw_text_mismatches": raw_mismatches,
        "imported_resume_cells_over_excel_limit": over_limit,
        "table_ref": worksheet.tables["HireXCandidates"].ref if "HireXCandidates" in worksheet.tables else "",
        "freeze_panes": str(worksheet.freeze_panes),
    }
    assert headers == FINAL_COLUMNS
    assert len(rows) == existing_count + appended_count
    assert len(headers) == 43
    assert blank_required_cells == 0
    assert duplicate_emails == 0
    assert duplicate_phones == 0
    assert duplicate_ids == 0
    assert raw_mismatches == 0
    assert over_limit == 0
    workbook.close()
    return checks


def main() -> None:
    random.seed(84)
    selected, inventory_stats = inventory_unique_resumes()
    results = append_rows_to_workbook(selected)
    print("Resume inventory:")
    for key, value in inventory_stats.items():
        print(f"{key}: {value}")
    print("Workbook results:")
    for key, value in results.items():
        print(f"{key}: {value}")
    print(f"Workbook updated: {WORKBOOK_PATH}")


if __name__ == "__main__":
    main()
